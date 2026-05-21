"""Create a Word-friendly notation handout for the poster."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


THIS_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = THIS_DIR / "evaluation_framework_notation.docx"


def add_equation(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    run = paragraph.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(10.5)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.05
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Evaluation Framework: Notation for Revealed-Behavior Matching")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(14)

    add_body(
        document,
        "This page defines the notation used for the behavioral skewness plot. "
        "The equations are written in a Word/PowerPoint equation-editor style.",
    )

    document.add_heading("Objects", level=2)
    add_body(
        document,
        "Let g in G index target social-interaction games, and let I_g be the set "
        "of real human players observed in game g.",
    )
    add_equation(document, "I_g = {1, ..., n_g}")
    add_body(
        document,
        "Each player has a revealed behavior trajectory b_{gi} and a pre-specified "
        "behavioral feature vector x_{gi}.",
    )
    add_equation(document, "x_{gi} = f(b_{gi}) in R^d")

    document.add_heading("Human Reference Distribution", level=2)
    add_body(
        document,
        "The empirical human reference distribution is uniform over the observed "
        "player trajectories within the same game.",
    )
    add_equation(document, "P_g(i) = 1 / n_g,    i in I_g")

    document.add_heading("LLM-Matched Distribution", level=2)
    add_body(
        document,
        "For persona source s, let a = 1, ..., m_s index sampled personas. Given "
        "persona a and the full transcript of game g, the LLM returns a top-K "
        "probability distribution over players. Unlisted players receive probability 0; "
        "in our experiments K = 3.",
    )
    add_equation(document, "r_{gasi} >= 0,    sum_{i in I_g} r_{gasi} = 1")
    add_body(document, "Aggregating over personas gives the matched distribution:")
    add_equation(document, "Q_{gs}(i) = (1 / m_s) sum_{a=1}^{m_s} r_{gasi}")
    add_body(document, "For the no-persona baseline, m_s = 1.")

    document.add_heading("Behavioral Skewness", level=2)
    add_body(document, "For behavioral feature l, define the human and matched means:")
    add_equation(document, "mu^P_{gl}  = sum_{i in I_g} P_g(i) x_{gil}")
    add_equation(document, "mu^Q_{gsl} = sum_{i in I_g} Q_{gs}(i) x_{gil}")
    add_body(document, "The game-level behavioral skew is matched minus human:")
    add_equation(document, "Delta_{gsl} = mu^Q_{gsl} - mu^P_{gl}")
    add_body(
        document,
        "The plotted statistic averages across games and standardizes by the "
        "empirical human standard deviation for that feature:",
    )
    add_equation(
        document,
        "Delta_tilde_{sl} = [(1 / |G|) sum_{g in G} Delta_{gsl}] / sigma^P_l",
    )
    add_equation(document, "sigma^P_l = SD_{g in G, i drawn from P_g}(x_{gil})")

    document.add_heading("Interpretation", level=2)
    add_bullet(document, "Delta_tilde_{sl} = 0: no average skew relative to real human trajectories.")
    add_bullet(
        document,
        "Delta_tilde_{sl} > 0: persona source s over-selects players with higher values of behavior l.",
    )
    add_bullet(
        document,
        "Delta_tilde_{sl} < 0: persona source s under-selects players with higher values of behavior l.",
    )
    add_bullet(
        document,
        "If a persona source spans the target behavior distribution, Q_{gs} should be close to P_g across behavioral features.",
    )

    document.add_heading("Caption Draft", level=2)
    add_body(
        document,
        "Behavioral skewness of persona-conditioned matches. For each target game, "
        "the empirical human reference distribution P_g is uniform over all real "
        "players observed in that game. For each persona source s, the matched "
        "distribution Q_{gs} is obtained by averaging the LLM's top-3 probability "
        "distributions over sampled personas. Each cell shows the matched-minus-human "
        "difference in a behavioral feature, standardized by the empirical human "
        "standard deviation of that feature. Values near zero indicate that the "
        "selected trajectories match the human reference distribution on that feature; "
        "positive and negative values indicate over-selection and under-selection, respectively.",
    )

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
