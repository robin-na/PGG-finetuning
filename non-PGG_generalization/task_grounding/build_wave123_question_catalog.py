#!/usr/bin/env python3
"""Export all catalog questions with full text and wave-1/2/3 presence flags."""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Dict, Iterable, List, Optional

from datasets import load_dataset, load_from_disk
from docx import Document
from huggingface_hub import hf_hub_download


REPO_ID = "LLM-Digital-Twin/Twin-2K-500"
CONFIG = "wave_split"
QUESTION_CATALOG_FILE = "question_catalog_and_human_response_csv/question_catalog.json"
OUT_DIR = Path("non-PGG_generalization/task_grounding")
LOCAL_SNAPSHOT_ROOT = Path("non-PGG_generalization/data/Twin-2k-500/snapshot")
LOCAL_QUESTION_CATALOG = LOCAL_SNAPSHOT_ROOT / QUESTION_CATALOG_FILE
LOCAL_WAVE_SPLIT = Path("non-PGG_generalization/data/Twin-2k-500/wave_split_dataset")
QUESTIONNAIRE_DIR = LOCAL_SNAPSHOT_ROOT / "raw_data" / "questionnaire"


PERSONA_BLOCK_TO_WAVE = {
    0: 1,
    1: 1,
    2: 1,
    3: 1,
    4: 1,
    5: 2,
    6: 2,
    7: 2,
    8: 2,
    9: 2,
    10: 3,
    11: 3,
    12: 3,
}


def load_question_catalog() -> List[Dict]:
    try:
        catalog_path = hf_hub_download(
            repo_id=REPO_ID,
            filename=QUESTION_CATALOG_FILE,
            repo_type="dataset",
        )
    except Exception:
        if not LOCAL_QUESTION_CATALOG.exists():
            raise
        catalog_path = str(LOCAL_QUESTION_CATALOG)
    with Path(catalog_path).open("r", encoding="utf-8") as f:
        return json.load(f)


def load_wave_split():
    try:
        return load_dataset(REPO_ID, CONFIG)["data"]
    except Exception:
        if not LOCAL_WAVE_SPLIT.exists():
            raise
        return load_from_disk(str(LOCAL_WAVE_SPLIT))["data"]


def normalize(text: str) -> str:
    return " ".join((text or "").split()).lower()


def docx_text(path: Path) -> str:
    doc = Document(str(path))
    chunks: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    chunks.append(text)
    return normalize("\n".join(chunks))


def build_wave_docs() -> Dict[int, str]:
    out = {}
    for wave in (1, 2, 3):
        path = QUESTIONNAIRE_DIR / f"Digital_Twins_-_Wave_{wave} with flow.docx"
        out[wave] = docx_text(path)
    return out


def block_wave_lookup(first_example) -> Dict[str, int]:
    blocks = json.loads(first_example["wave1_3_persona_json"])
    ref_to_wave: Dict[str, int] = {}
    for idx, block in enumerate(blocks):
        wave = PERSONA_BLOCK_TO_WAVE.get(idx)
        if wave is None:
            continue
        block_name = block.get("BlockName", "")
        for q in block.get("Questions", []):
            ref = f"{block_name}::{q.get('QuestionID', '')}"
            ref_to_wave[ref] = wave
    return ref_to_wave


def search_fragments(question: Dict) -> List[str]:
    fragments: List[str] = []
    question_text = normalize(question.get("QuestionText", ""))
    if question_text:
        fragments.append(question_text)

    rows = question.get("Rows", []) or []
    for row in rows[:2]:
        row_text = normalize(str(row))
        if row_text:
            fragments.append(row_text)
            break

    options = question.get("Options", []) or []
    for option in options[:1]:
        opt_text = normalize(str(option))
        if opt_text:
            fragments.append(opt_text)
            break

    deduped: List[str] = []
    seen = set()
    for fragment in fragments:
        if fragment and fragment not in seen:
            deduped.append(fragment)
            seen.add(fragment)
    return deduped


def wave_presence_from_docs(question: Dict, wave_docs: Dict[int, str]) -> Dict[int, bool]:
    fragments = search_fragments(question)
    presence = {1: False, 2: False, 3: False}
    for wave, doc_text in wave_docs.items():
        if fragments and all(fragment in doc_text for fragment in fragments):
            presence[wave] = True
    return presence


def build_rows(
    catalog: List[Dict],
    ref_to_wave: Dict[str, int],
    wave_docs: Dict[int, str],
) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for question in catalog:
        ref = f"{question.get('BlockName', '')}::{question.get('QuestionID', '')}"
        source = question.get("source", "")
        if source == "wave1_3_persona_json":
            wave = ref_to_wave.get(ref)
            presence = {1: wave == 1, 2: wave == 2, 3: wave == 3}
            method = "persona_block_order"
        else:
            presence = wave_presence_from_docs(question, wave_docs)
            method = "questionnaire_text_match"

        rows.append(
            {
                "question_id": question.get("QuestionID", ""),
                "block_name": question.get("BlockName", ""),
                "question_type": question.get("QuestionType", ""),
                "source": source,
                "wave_1_exists": str(bool(presence[1])),
                "wave_2_exists": str(bool(presence[2])),
                "wave_3_exists": str(bool(presence[3])),
                "wave_presence": ",".join(
                    str(wave)
                    for wave in (1, 2, 3)
                    if presence[wave]
                ),
                "wave_mapping_method": method,
                "question_text_full": " ".join((question.get("QuestionText", "") or "").split()),
                "options_json": json.dumps(question.get("Options", []), ensure_ascii=False),
                "rows_json": json.dumps(question.get("Rows", []), ensure_ascii=False),
                "csv_columns_json": json.dumps(question.get("csv_columns", []), ensure_ascii=False),
            }
        )
    rows.sort(key=lambda r: (r["block_name"], r["question_id"]))
    return rows


def write_csv(path: Path, rows: Iterable[Dict[str, str]]) -> None:
    rows = list(rows)
    if not rows:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def main() -> int:
    catalog = load_question_catalog()
    ds = load_wave_split()
    ref_to_wave = block_wave_lookup(ds[0])
    wave_docs = build_wave_docs()
    rows = build_rows(catalog, ref_to_wave, wave_docs)
    out_path = OUT_DIR / "wave123_question_catalog.csv"
    write_csv(out_path, rows)
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
